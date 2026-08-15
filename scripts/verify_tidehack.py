from pathlib import Path
from zipfile import ZipFile
from docx import Document
base=Path(r'C:\Users\ASUS\Documents\官网\docs')
for p in [base/'青岛潮汐TIDEHACK2026_完整策划书.docx', base/'青岛潮汐TIDEHACK2026_对外合作方案.docx']:
    with ZipFile(p) as z:
        bad=z.testzip()
        assert bad is None, bad
    d=Document(p)
    headings=[x.text for x in d.paragraphs if x.style.name.startswith('Heading')]
    tables=len(d.tables)
    text='\n'.join(x.text for x in d.paragraphs)
    print(p.name, 'bytes=',p.stat().st_size, 'paragraphs=',len(d.paragraphs), 'tables=',tables, 'headings=',len(headings), 'has_tide=', 'TIDEHACK' in text)
    print(' first headings:', headings[:5])
print('QA internal pages:',len(list((base/'qa_internal_v5').glob('page-*.png'))))
print('QA external pages:',len(list((base/'qa_external_v5').glob('page-*.png'))))
