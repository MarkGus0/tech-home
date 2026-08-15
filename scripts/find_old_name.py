from docx import Document
from pathlib import Path
for p in [Path(r'C:\Users\ASUS\Documents\官网\docs\未完场UNFINO2026_完整策划书.docx'), Path(r'C:\Users\ASUS\Documents\官网\docs\未完场UNFINO2026_对外合作方案.docx')]:
    d=Document(p); print('\n',p.name)
    for x in d.paragraphs:
        if 'TIDEHACK' in x.text or 'tidehack' in x.text.lower(): print('P:',repr(x.text))
    for ti,t in enumerate(d.tables):
        for ri,r in enumerate(t.rows):
            for ci,c in enumerate(r.cells):
                if 'TIDEHACK' in c.text or 'tidehack' in c.text.lower(): print('T',ti,ri,ci,repr(c.text))
