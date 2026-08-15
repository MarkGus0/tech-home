from pathlib import Path
from docx import Document
base=Path(r'C:\Users\ASUS\Documents\官网\docs')
brand='TechFlows 社区旗下子品牌「未完场 UNFINO」'
for p in [base/'青岛潮汐TIDEHACK2026_完整策划书.docx', base/'青岛潮汐TIDEHACK2026_对外合作方案.docx']:
    d=Document(p)
    text='\n'.join(x.text for x in d.paragraphs)
    text+='\n'+'\n'.join(c.text for t in d.tables for r in t.rows for c in r.cells)
    print(p.name, 'brand_count=', text.count(brand), 'unfino_count=', text.count('UNFINO'), 'bytes=', p.stat().st_size)
    assert 'UNFINO' in text
