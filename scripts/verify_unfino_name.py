from pathlib import Path
from docx import Document
base=Path(r'C:\Users\ASUS\Documents\官网\docs')
files=[base/'未完场UNFINO2026_完整策划书.docx', base/'未完场UNFINO2026_对外合作方案.docx']
for p in files:
    d=Document(p)
    text='\n'.join(x.text for x in d.paragraphs)+'\n'+'\n'.join(c.text for t in d.tables for r in t.rows for c in r.cells)
    print(p.name, 'exists=',p.exists(),'bytes=',p.stat().st_size,'UNFINO 2026=',text.count('未完场 UNFINO 2026'),'old_tidehack=',text.count('TIDEHACK'))
    assert '未完场 UNFINO 2026' in text
    assert 'TIDEHACK' not in text
