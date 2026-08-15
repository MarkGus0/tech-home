from pathlib import Path
from docx import Document
base=Path(r'C:\Users\ASUS\Documents\官网\docs')
for p in [base/'未完场UNFINO2026_完整策划书.docx',base/'未完场UNFINO2026_对外合作方案.docx']:
 d=Document(p); text='\n'.join(x.text for x in d.paragraphs)+'\n'+'\n'.join(c.text for t in d.tables for r in t.rows for c in r.cells)
 print(p.name, 'UNFINO 2026=',text.count('未完场 UNFINO 2026'),'old_tidehack=',text.count('TIDEHACK'),'logo_unfino=',text.count('UNFINO'))
 assert '未完场 UNFINO 2026' in text and 'TIDEHACK' not in text
